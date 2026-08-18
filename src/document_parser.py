# src/document_parser.py

import fitz  # PyMuPDF
import pdfplumber
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import io
import re
import logging
import numpy as np
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import json
import hashlib
from collections import defaultdict

logger = logging.getLogger(__name__)

# Проверка дополнительных библиотек
try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False
    logger.warning("EasyOCR не установлен")

try:
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    CAMELOT_AVAILABLE = False
    logger.warning("Camelot не установлен")


class DocumentParser:
    """
    Продвинутый парсер документов с поддержкой:
    - Многоуровневого извлечения текста с сохранением структуры
    - Распознавания таблиц с помощью нескольких методов
    - OCR для изображений и скриншотов
    - Сохранения позиционирования элементов
    - Извлечения метаданных и структуры документа
    - Кэширование результатов в cache
    """
    
    def __init__(self, use_ocr: bool = True, ocr_lang: str = 'rus+eng', cache_dir: Path = None):
        self.use_ocr = use_ocr
        self.ocr_lang = ocr_lang
        self.tesseract_available = self._check_tesseract()
        self.easyocr_available = EASYOCR_AVAILABLE
        self.easyocr_reader = None
        
        # Настройка кэша
        try:
            from src.config import Config
            self.cache_dir = cache_dir or Config.PARSED_DOCS_DIR
        except ImportError:
            self.cache_dir = cache_dir or Path("cache/parsed_documents")
        
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        # Кэш результатов
        self._cache = {}
        self._load_cache()
        
        if self.easyocr_available and self.use_ocr:
            try:
                lang_list = ['en']
                if 'rus' in self.ocr_lang or 'ru' in self.ocr_lang:
                    lang_list.append('ru')
                self.easyocr_reader = easyocr.Reader(lang_list, gpu=False, verbose=False)
                logger.info(f"✅ EasyOCR инициализирован: {lang_list}")
            except Exception as e:
                logger.warning(f"Ошибка инициализации EasyOCR: {e}")
                self.easyocr_available = False
    
    def _check_tesseract(self) -> bool:
        """Проверка доступности Tesseract"""
        try:
            pytesseract.get_tesseract_version()
            return True
        except:
            return False
    
    def _get_file_hash(self, file_path: Path) -> str:
        """Вычисляет хеш файла для кэширования"""
        hasher = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                hasher.update(chunk)
        return hasher.hexdigest()
    
    def _get_cache_path(self, file_path: Path) -> Path:
        """Возвращает путь к кэш-файлу для документа"""
        file_hash = self._get_file_hash(file_path)
        return self.cache_dir / f"{file_path.stem}_{file_hash[:12]}.json"
    
    def _load_cache(self):
        """Загружает кэш из файлов"""
        for cache_file in self.cache_dir.glob("*.json"):
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    key = cache_file.stem
                    self._cache[key] = data
                    logger.info(f"📦 Загружен кэш: {cache_file.name}")
            except Exception as e:
                logger.warning(f"⚠️ Ошибка загрузки кэша {cache_file}: {e}")
    
    def _save_to_cache(self, file_path: Path, result: Dict[str, Any]):
        """Сохраняет результат в кэш"""
        cache_path = self._get_cache_path(file_path)
        try:
            clean_result = self._prepare_for_cache(result)
            with open(cache_path, 'w', encoding='utf-8') as f:
                json.dump(clean_result, f, ensure_ascii=False, indent=2)
            logger.info(f"💾 Сохранен кэш: {cache_path.name}")
        except Exception as e:
            logger.warning(f"⚠️ Ошибка сохранения кэша: {e}")
    
    def _prepare_for_cache(self, result: Dict[str, Any]) -> Dict[str, Any]:
        """Подготавливает результат для сохранения в JSON"""
        clean = {}
        
        for key, value in result.items():
            if key == 'pages':
                clean_pages = []
                for page in value:
                    clean_page = {k: v for k, v in page.items() 
                                if k not in ['raw_blocks', 'image']}
                    if 'images' in clean_page:
                        clean_page['images'] = [
                            {k: v for k, v in img.items() 
                             if k not in ['image', 'image_data']}
                            for img in clean_page['images']
                        ]
                    clean_pages.append(clean_page)
                clean['pages'] = clean_pages
            elif key == 'images':
                clean['images'] = [
                    {k: v for k, v in img.items() 
                     if k not in ['image', 'image_data']}
                    for img in value
                ]
            else:
                clean[key] = value
        
        return clean
    
    def parse_document(self, file_path: Path) -> Dict[str, Any]:
        """Парсинг документа с использованием кэша"""
        # Проверяем кэш
        cache_path = self._get_cache_path(file_path)
        if cache_path.exists():
            try:
                with open(cache_path, 'r', encoding='utf-8') as f:
                    result = json.load(f)
                logger.info(f"📦 Использован кэш: {cache_path.name}")
                return result
            except Exception as e:
                logger.warning(f"⚠️ Ошибка загрузки кэша: {e}")
        
        # Парсим документ
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            result = self._parse_pdf(file_path)
        elif ext == '.docx':
            result = self._parse_docx(file_path)
        elif ext == '.doc':
            result = self._parse_doc(file_path)
        elif ext in ['.txt', '.md']:
            result = self._parse_txt(file_path)
        else:
            logger.warning(f"Неподдерживаемый формат: {ext}")
            result = self._empty_result(file_path, f"Unsupported format: {ext}")
        
        # Сохраняем в кэш
        self._save_to_cache(file_path, result)
        
        return result
    
    def parse_all_documents(self, documents_dir: Path) -> List[Dict[str, Any]]:
        """Парсинг всех документов в директории"""
        results = []
        extensions = ['.pdf', '.docx', '.doc', '.txt', '.md']
        
        for ext in extensions:
            for file_path in documents_dir.glob(f'*{ext}'):
                if file_path.is_file() and file_path.name != '.gitkeep':
                    logger.info(f"📄 Парсинг: {file_path.name}")
                    result = self.parse_document(file_path)
                    if result and result.get('pages'):
                        results.append(result)
                        logger.info(f"  ✅ {file_path.name}: {len(result['pages'])} страниц")
        
        return results
    
    def _empty_result(self, file_path: Path, error: str = "") -> Dict[str, Any]:
        """Возвращает пустой результат"""
        return {
            'filename': file_path.name,
            'full_text': '',
            'pages': [],
            'tables': [],
            'images': [],
            'headers': [],
            'structure': {'sections': [], 'appendices': []},
            'metadata': {
                'total_pages': 0,
                'has_tables': False,
                'has_images': False,
                'has_headers': False,
                'error': error
            }
        }
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Парсинг PDF файла"""
        logger.info(f"📄 Начинаем парсинг PDF: {file_path.name}")
        
        result = {
            'filename': file_path.name,
            'pages': [],
            'tables': [],
            'images': [],
            'headers': [],
            'structure': {'sections': [], 'appendices': []},
            'metadata': {
                'total_pages': 0,
                'has_tables': False,
                'has_images': False,
                'has_headers': False,
                'extraction_methods': []
            },
            'full_text': ''
        }
        
        try:
            doc = fitz.open(str(file_path))
            result['metadata']['total_pages'] = len(doc)
            
            all_text_parts = []
            
            for page_num in range(len(doc)):
                logger.info(f"  📝 Страница {page_num + 1}/{len(doc)}")
                
                page_data = self._parse_page(doc, page_num, file_path)
                result['pages'].append(page_data)
                
                if page_data.get('text'):
                    all_text_parts.append(f"[Страница {page_num + 1}]\n{page_data['text']}")
                
                if page_data.get('tables'):
                    for table in page_data['tables']:
                        table['page'] = page_num + 1
                        result['tables'].append(table)
                    result['metadata']['has_tables'] = True
                
                if page_data.get('images'):
                    for img in page_data['images']:
                        img['page'] = page_num + 1
                        result['images'].append(img)
                    result['metadata']['has_images'] = True
                
                if page_data.get('headers'):
                    for header in page_data['headers']:
                        header['page'] = page_num + 1
                        result['headers'].append(header)
                    result['metadata']['has_headers'] = True
            
            doc.close()
            
            result['full_text'] = '\n\n'.join(all_text_parts)
            result['structure'] = self._analyze_structure(result['pages'])
            
            logger.info(f"✅ PDF спарсен: {len(result['pages'])} страниц, "
                       f"{len(result['tables'])} таблиц, {len(result['images'])} изображений")
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Ошибка парсинга PDF: {e}", exc_info=True)
            return self._empty_result(file_path, str(e))
    
    def _parse_page(self, doc: fitz.Document, page_num: int, file_path: Path) -> Dict[str, Any]:
        """Парсинг отдельной страницы"""
        page = doc[page_num]
        page_data = {
            'page_num': page_num + 1,
            'type': 'text',
            'text': '',
            'headers': [],
            'tables': [],
            'images': [],
            'has_tables': False,
            'has_images': False,
            'blocks': []
        }
        
        # 1. Извлекаем текст
        page_text = page.get_text()
        page_data['text'] = page_text
        
        # 2. Определяем тип страницы
        page_data['type'] = self._detect_page_type(page, page_text)
        
        # 3. Извлекаем заголовки
        headers = self._extract_headers(page_text)
        if headers:
            page_data['headers'] = headers
        
        # 4. Извлекаем таблицы
        tables = self._extract_tables(page, page_text)
        if tables:
            page_data['tables'] = tables
            page_data['has_tables'] = True
        
        # 5. Извлекаем изображения и применяем OCR
        if self.use_ocr:
            images = self._extract_images(page, page_num + 1)
            if images:
                page_data['images'] = images
                page_data['has_images'] = True
                
                ocr_texts = self._perform_ocr_on_images(images)
                if ocr_texts:
                    page_data['ocr_text'] = '\n\n'.join(ocr_texts)
                    page_data['text'] += '\n\n[OCR TEXT]\n' + page_data['ocr_text']
        
        return page_data
    
    def _detect_page_type(self, page: fitz.Page, text: str) -> str:
        """Определение типа страницы"""
        image_list = page.get_images(full=True)
        has_images = len(image_list) > 0
        has_text = bool(text.strip())
        
        has_tables = self._visual_table_detection(page, text)
        
        if has_images and not has_text:
            return 'image'
        elif has_tables and has_text:
            return 'mixed'
        elif has_tables:
            return 'table'
        elif has_images:
            return 'image'
        else:
            return 'text'
    
    def _visual_table_detection(self, page: fitz.Page, text: str) -> bool:
        """Визуальное обнаружение таблиц"""
        try:
            # Проверяем наличие линий
            draws = page.get_drawings()
            h_lines = 0
            v_lines = 0
            
            for draw in draws:
                if draw.get('type') == 'l':
                    rect = draw.get('rect', (0, 0, 0, 0))
                    width = rect[2] - rect[0]
                    height = rect[3] - rect[1]
                    
                    if height < 5 and width > 50:
                        h_lines += 1
                    elif width < 5 and height > 50:
                        v_lines += 1
            
            if h_lines >= 2 and v_lines >= 2:
                return True
            
            # Проверяем текст
            lines = text.split('\n')
            table_indicators = 0
            
            for line in lines:
                if '|' in line or '\t' in line:
                    table_indicators += 1
                if len(line.strip()) > 20 and line.count(' ') > 3:
                    table_indicators += 0.5
            
            return table_indicators >= 3
            
        except Exception as e:
            logger.warning(f"Ошибка визуального обнаружения таблиц: {e}")
            return False
    
    def _extract_headers(self, text: str) -> List[Dict[str, Any]]:
        """Извлечение заголовков"""
        headers = []
        
        patterns = [
            (r'^(\d+\.\d+\.\d+\s+[А-ЯЁ][а-яё\s\d]+)', 4),
            (r'^(\d+\.\d+\s+[А-ЯЁ][а-яё\s\d]+)', 3),
            (r'^(\d+\s+[А-ЯЁ][а-яё\s\d]+)', 2),
            (r'^([А-ЯЁ][А-ЯЁ\s\d.]+[А-ЯЁ])$', 1),
            (r'^(Приложение\s+\d+[\.\s]*[А-ЯЁа-яё\s\d]+)', 1),
            (r'^(Требования\s+к\s+[А-ЯЁа-яё\s]+)', 2),
            (r'^(Порядок\s+[А-ЯЁа-яё\s]+)', 2),
            (r'^(Раздел\s+\d+[\.\s]*[А-ЯЁа-яё\s\d]+)', 2),
            (r'^(Глава\s+\d+[\.\s]*[А-ЯЁа-яё\s\d]+)', 2),
        ]
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line or len(line) < 5 or len(line) > 200:
                continue
            
            for pattern, level in patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    header_text = match.group(1).strip()
                    if not self._is_false_header(header_text):
                        headers.append({
                            'text': header_text,
                            'level': level,
                            'pattern': pattern
                        })
                    break
        
        return headers
    
    def _is_false_header(self, text: str) -> bool:
        """Проверка ложного заголовка"""
        if len(text) < 5:
            return True
        if text.endswith('.') and not text.isupper():
            return True
        if re.match(r'^\d+$', text):
            return True
        false_patterns = [
            r'^ООО\s+',
            r'^УНП\s+',
            r'^Контактный',
            r'^Телефон',
            r'^Факс',
            r'^Адрес',
            r'^E-mail',
            r'^http',
            r'^www',
            r'^\d+\.\d+\s+руб',
        ]
        for pattern in false_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        return False
    
    def _extract_tables(self, page: fitz.Page, text: str) -> List[Dict[str, Any]]:
        """Извлечение таблиц"""
        tables = []
        
        try:
            # PyMuPDF find_tables
            found_tables = page.find_tables()
            if found_tables and found_tables.tables:
                for idx, table in enumerate(found_tables.tables, 1):
                    table_data = table.extract()
                    if table_data and len(table_data) > 1:
                        cleaned = self._clean_table_data(table_data)
                        if cleaned:
                            tables.append({
                                'id': idx,
                                'source': 'pymupdf',
                                'data': cleaned,
                                'text_representation': self._format_table(cleaned)
                            })
            
            # Текстовые таблицы
            text_tables = self._extract_text_tables(text)
            if text_tables:
                tables.extend(text_tables)
                
        except Exception as e:
            logger.warning(f"Ошибка извлечения таблиц PyMuPDF: {e}")
        
        return tables
    
    def _extract_text_tables(self, text: str) -> List[Dict[str, Any]]:
        """Извлечение таблиц из текста"""
        tables = []
        lines = text.split('\n')
        
        table_blocks = []
        current_block = []
        
        for line in lines:
            is_table_line = (
                '|' in line or 
                '\t' in line or 
                (line.count(' ') > 5 and len(line) > 30) or
                (line.count(' ') > 3 and len(line) > 50 and any(c.isdigit() for c in line))
            )
            
            if is_table_line:
                current_block.append(line)
            elif current_block and len(current_block) > 2:
                table_blocks.append(current_block)
                current_block = []
        
        if current_block and len(current_block) > 2:
            table_blocks.append(current_block)
        
        for block_idx, block in enumerate(table_blocks, 1):
            table_data = []
            for line in block:
                if '|' in line:
                    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                elif '\t' in line:
                    cells = [cell.strip() for cell in line.split('\t') if cell.strip()]
                else:
                    cells = re.split(r'\s{2,}', line.strip())
                    cells = [c.strip() for c in cells if c.strip()]
                
                if cells and len(cells) > 1:
                    table_data.append(cells)
            
            if table_data and len(table_data) > 1:
                tables.append({
                    'id': f'text_{block_idx}',
                    'source': 'text_extraction',
                    'data': table_data,
                    'text_representation': self._format_table(table_data)
                })
        
        return tables
    
    def _clean_table_data(self, table_data: List[List[str]]) -> List[List[str]]:
        """Очистка данных таблицы"""
        cleaned = []
        for row in table_data:
            clean_row = [str(cell).strip() if cell else '' for cell in row]
            if any(cell for cell in clean_row):
                cleaned.append(clean_row)
        return cleaned
    
    def _format_table(self, table_data: List[List[str]]) -> str:
        """Форматирование таблицы"""
        if not table_data:
            return ""
        
        col_count = max(len(row) for row in table_data)
        col_widths = []
        
        for col_idx in range(col_count):
            max_width = 0
            for row in table_data:
                if col_idx < len(row):
                    max_width = max(max_width, len(str(row[col_idx])))
            col_widths.append(max_width + 2)
        
        lines = []
        for row_idx, row in enumerate(table_data):
            padded_row = row + [''] * (col_count - len(row))
            row_str = "| "
            for col_idx, cell in enumerate(padded_row):
                row_str += str(cell).ljust(col_widths[col_idx]) + " | "
            lines.append(row_str.strip())
            
            if row_idx == 0 and len(table_data) > 1:
                separator = "+" + "+".join(['-' * (w + 1) for w in col_widths]) + "+"
                lines.append(separator)
        
        return "\n".join(lines)
    
    def _extract_images(self, page: fitz.Page, page_num: int) -> List[Dict[str, Any]]:
        """Извлечение изображений"""
        images = []
        image_list = page.get_images(full=True)
        
        # Получаем документ из страницы
        doc = page.parent
        
        for img_idx, img in enumerate(image_list, 1):
            try:
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                
                if pix.n - pix.alpha < 4:
                    img_data = pix.tobytes("png")
                    pil_image = Image.open(io.BytesIO(img_data))
                else:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                    img_data = pix.tobytes("png")
                    pil_image = Image.open(io.BytesIO(img_data))
                
                width, height = pil_image.size
                
                images.append({
                    'id': f'page_{page_num}_img_{img_idx}',
                    'page': page_num,
                    'xref': xref,
                    'width': width,
                    'height': height,
                    'image': pil_image,
                    'image_data': img_data,
                    'has_text': False,
                    'ocr_text': ''
                })
                
                pix = None
                
            except Exception as e:
                logger.warning(f"Ошибка извлечения изображения {img_idx}: {e}")
        
        return images
    
    def _perform_ocr_on_images(self, images: List[Dict[str, Any]]) -> List[str]:
        """Применение OCR к изображениям"""
        ocr_results = []
        
        for img_info in images:
            try:
                image = img_info['image']
                processed = self._preprocess_image(image)
                text = self._perform_hybrid_ocr(processed)
                
                if text and len(text.strip()) > 10:
                    img_info['has_text'] = True
                    img_info['ocr_text'] = text
                    ocr_results.append(f"[Изображение {img_info['id']}]\n{text}")
                    logger.info(f"  ✅ OCR: {len(text)} символов")
                else:
                    logger.warning(f"  ⚠️ OCR: текст не распознан")
                    
            except Exception as e:
                logger.warning(f"❌ Ошибка OCR: {e}")
        
        return ocr_results
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """Предобработка изображения"""
        try:
            if image.mode != 'L':
                image = image.convert('L')
            
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            image = image.filter(ImageFilter.SHARPEN)
            image = image.point(lambda x: 0 if x < 128 else 255, '1')
            image = image.convert('L')
            return image
        except Exception as e:
            logger.warning(f"Ошибка предобработки: {e}")
            return image
    
    def _perform_hybrid_ocr(self, image: Image.Image) -> str:
        """Гибридный OCR"""
        results = {}
        
        if self.tesseract_available:
            try:
                custom_config = f'--oem 3 --psm 6 -l {self.ocr_lang}'
                text = pytesseract.image_to_string(image, config=custom_config)
                if text.strip():
                    results['tesseract'] = text
            except Exception as e:
                pass
        
        if self.easyocr_available and self.easyocr_reader:
            try:
                image_np = np.array(image)
                result = self.easyocr_reader.readtext(image_np, detail=0)
                text = ' '.join(result)
                if text.strip():
                    results['easyocr'] = text
            except Exception as e:
                pass
        
        best_text = self._select_best_ocr_result(results)
        return self._clean_ocr_text(best_text)
    
    def _select_best_ocr_result(self, results: Dict[str, str]) -> str:
        """Выбор лучшего результата OCR"""
        if not results:
            return ""
        
        if len(results) == 1:
            return list(results.values())[0]
        
        best_text = results.get('tesseract', '')
        easyocr_text = results.get('easyocr', '')
        
        if easyocr_text and len(easyocr_text) > len(best_text) * 0.8:
            best_ratio = self._text_readability_ratio(best_text)
            easy_ratio = self._text_readability_ratio(easyocr_text)
            if easy_ratio > best_ratio:
                best_text = easyocr_text
        
        return best_text
    
    def _text_readability_ratio(self, text: str) -> float:
        """Вычисление соотношения читаемых символов"""
        if not text:
            return 0.0
        readable = sum(1 for c in text if c.isalnum() or c in ' .,!?-')
        return readable / len(text) if text else 0.0
    
    def _clean_ocr_text(self, text: str) -> str:
        """Очистка OCR текста"""
        if not text:
            return ""
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            readability = self._text_readability_ratio(line)
            if readability > 0.4 and len(line) > 3:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _analyze_structure(self, pages: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Анализ структуры документа"""
        structure = {
            'sections': [],
            'appendices': [],
            'has_table_of_contents': False,
            'total_pages': len(pages)
        }
        
        all_headers = []
        for page in pages:
            for header in page.get('headers', []):
                all_headers.append({
                    'text': header['text'],
                    'level': header['level'],
                    'page': page['page_num']
                })
        
        if all_headers:
            for header in all_headers:
                text = header['text'].lower()
                if any(part in text for part in ['приложение', 'appendix']):
                    structure['appendices'].append(header)
                else:
                    structure['sections'].append(header)
        
        for page in pages[:3]:
            text = page.get('text', '').lower()
            if 'оглавление' in text or 'содержание' in text:
                structure['has_table_of_contents'] = True
                break
        
        return structure
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Парсинг DOCX файла"""
        try:
            import docx
            doc = docx.Document(str(file_path))
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            full_text = '\n'.join(text_parts)
            
            tables = []
            for table_idx, table in enumerate(doc.tables, 1):
                if table.rows:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_data.append(cell.text.strip())
                        if row_data:
                            table_data.append(row_data)
                    
                    if table_data and len(table_data) > 1:
                        tables.append({
                            'id': table_idx,
                            'source': 'docx',
                            'data': table_data,
                            'text_representation': self._format_table(table_data)
                        })
            
            headers = self._extract_headers(full_text)
            
            return {
                'filename': file_path.name,
                'pages': [{
                    'page_num': 1,
                    'text': full_text,
                    'type': 'text',
                    'headers': headers,
                    'tables': tables,
                    'has_tables': bool(tables),
                    'has_images': False,
                    'images': []
                }],
                'full_text': full_text,
                'tables': tables,
                'images': [],
                'headers': headers,
                'structure': self._analyze_structure([{
                    'page_num': 1,
                    'headers': headers,
                    'text': full_text
                }]),
                'metadata': {
                    'total_pages': 1,
                    'has_tables': bool(tables),
                    'has_images': False,
                    'has_headers': bool(headers)
                }
            }
        except Exception as e:
            logger.error(f"Ошибка парсинга DOCX: {e}")
            return self._empty_result(file_path, str(e))
    
    def _parse_doc(self, file_path: Path) -> Dict[str, Any]:
        """Парсинг DOC файла"""
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            
            text_pattern = re.compile(rb'[\x20-\x7E\x0410-\x044F]{20,}')
            matches = text_pattern.findall(content)
            
            text_parts = []
            for match in matches:
                try:
                    text_parts.append(match.decode('utf-8', errors='ignore'))
                except:
                    try:
                        text_parts.append(match.decode('cp1251', errors='ignore'))
                    except:
                        text_parts.append(match.decode('latin-1', errors='ignore'))
            
            full_text = '\n'.join(text_parts)
            headers = self._extract_headers(full_text)
            
            return {
                'filename': file_path.name,
                'pages': [{
                    'page_num': 1,
                    'text': full_text,
                    'type': 'text',
                    'headers': headers,
                    'tables': [],
                    'has_tables': False,
                    'has_images': False,
                    'images': []
                }],
                'full_text': full_text,
                'tables': [],
                'images': [],
                'headers': headers,
                'structure': self._analyze_structure([{
                    'page_num': 1,
                    'headers': headers,
                    'text': full_text
                }]),
                'metadata': {
                    'total_pages': 1,
                    'has_tables': False,
                    'has_images': False,
                    'has_headers': bool(headers)
                }
            }
        except Exception as e:
            logger.error(f"Ошибка парсинга DOC: {e}")
            return self._empty_result(file_path, str(e))
    
    def _parse_txt(self, file_path: Path) -> Dict[str, Any]:
        """Парсинг TXT файла"""
        try:
            encodings = ['utf-8', 'cp1251', 'koi8-r', 'latin-1']
            full_text = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        full_text = f.read()
                        break
                except UnicodeDecodeError:
                    continue
            else:
                with open(file_path, 'rb') as f:
                    full_text = f.read().decode('utf-8', errors='ignore')
            
            headers = self._extract_headers(full_text)
            
            return {
                'filename': file_path.name,
                'pages': [{
                    'page_num': 1,
                    'text': full_text,
                    'type': 'text',
                    'headers': headers,
                    'tables': [],
                    'has_tables': False,
                    'has_images': False,
                    'images': []
                }],
                'full_text': full_text,
                'tables': [],
                'images': [],
                'headers': headers,
                'structure': self._analyze_structure([{
                    'page_num': 1,
                    'headers': headers,
                    'text': full_text
                }]),
                'metadata': {
                    'total_pages': 1,
                    'has_tables': False,
                    'has_images': False,
                    'has_headers': bool(headers)
                }
            }
        except Exception as e:
            logger.error(f"Ошибка парсинга TXT: {e}")
            return self._empty_result(file_path, str(e))


# Глобальный экземпляр парсера
document_parser = DocumentParser()